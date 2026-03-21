#!/usr/bin/env python3
"""Convert methods+results.txt to formatted .docx with images and tables."""
import re, os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = '/Users/yu/code/code2601/TY'
TXT = f'{BASE}/docs/methods+results.txt'
OUT = f'{BASE}/docs/methods+results.docx'

# Figure placeholder line numbers (1-indexed) -> (image_path, caption)
FIG_MAP = {
    61: (f'{BASE}/output/figures/pretrend_diagnostic.png',
         '图 1：事前趋势诊断：加入基线控制变量前后的比较（Shannon 熵，基期 = 2021）'),
    77: (f'{BASE}/output/figures/trend_by_tercile.png',
         '图 2：技能综合性的分组趋势（按 AI 暴露度三分位）'),
    95: (f'{BASE}/output/figures/event_study_entropy_cond_pt.png',
         '图 3：事件研究——Shannon 熵的逐年系数（条件平行趋势规格，基期 = 2021）'),
    392:(f'{BASE}/output/figures/event_study_entropy_cond_pt.png',
         '图 4：随机化推断置换分布'),
}

HEADINGS = {
    '研究方法': 1, '技能综合性的测量': 2, '识别策略': 2,
    '实证结果': 1, '综合化趋势的初步事实': 2, 'AI 暴露度的边际效应': 2,
    '事件研究': 2, '异质性分析': 2, '暴露度阈值效应': 3, '行业异质性': 3,
    '补充证据：技能描述的内容变化': 2, '稳健性检验': 2,
    'LDA 主题一览': 2, '处理期定义的稳健性检验': 2,
    'Rambachan & Roth (2023) 敏感性分析': 2,
}


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def style_cell(cell, bold=False, size=Pt(10)):
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = size
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.bold = bold


def is_separator(s):
    return bool(s) and all(c in '-+:= |' for c in s) and len(s) > 4


def is_table_line(s, prev_s='', next_s=''):
    if not s:
        return False
    if is_separator(s):
        return True
    if '|' in s and not s.startswith('['):
        return True
    if re.search(r'\S  {2,}\S', s) and len(s) < 250 and not s.startswith('注：') and not s.startswith('注:'):
        if is_separator(prev_s) or is_separator(next_s) or \
           re.search(r'\S  {2,}\S', prev_s) or re.search(r'\S  {2,}\S', next_s):
            return True
    return False


def parse_table_rows(table_lines):
    """Parse table text lines into list of row-lists."""
    rows = []
    for line in table_lines:
        if is_separator(line):
            continue
        if '|' in line:
            cells = [c.strip() for c in line.split('|')]
            # Remove empty first/last from leading/trailing |
            if cells and cells[0] == '':
                cells = cells[1:]
            if cells and cells[-1] == '':
                cells = cells[:-1]
        else:
            cells = [c.strip() for c in re.split(r'  {2,}', line)]
        cells = [c for c in cells if c is not None]
        if cells:
            rows.append(cells)
    return rows


def add_word_table(doc, table_lines):
    rows = parse_table_rows(table_lines)
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < max_cols:
            r.append('')

    tbl = doc.add_table(rows=len(rows), cols=max_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, text in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.text = text
            is_header = (i == 0)
            style_cell(cell, bold=is_header)
            if is_header:
                set_cell_shading(cell, 'D9E2F3')


def add_figure(doc, path, caption):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(5.5))
    else:
        p = doc.add_paragraph(f'[图片缺失: {path}]')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(10)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.italic = True


def main():
    doc = Document()

    # Default style
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading('生成式 AI 是否推动了岗位技能需求综合化？', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    with open(TXT, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_table = False
    table_buf = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            add_word_table(doc, table_buf)
            table_buf = []

    while i < len(lines):
        line = lines[i].rstrip('\n')
        ln = i + 1
        s = line.strip()

        prev_s = lines[i-1].strip() if i > 0 else ''
        next_s = lines[i+1].strip() if i+1 < len(lines) else ''

        # Figure placeholder
        if s == '[]' and ln in FIG_MAP:
            flush_table()
            in_table = False
            path, cap = FIG_MAP[ln]
            add_figure(doc, path, cap)
            i += 1
            continue

        # Heading
        if s in HEADINGS:
            flush_table()
            in_table = False
            h = doc.add_heading(s, level=HEADINGS[s])
            for run in h.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
            continue

        # Table caption (": xxx")
        if s.startswith(': ') and len(s) > 2:
            flush_table()
            in_table = False
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(s[2:])
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            i += 1
            continue

        # Table line detection
        if is_table_line(s, prev_s, next_s):
            in_table = True
            table_buf.append(s)
            i += 1
            continue
        elif in_table:
            flush_table()
            in_table = False
            # Don't increment, reprocess current line
            continue

        # Empty line
        if not s:
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph(s)
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        i += 1

    flush_table()
    doc.save(OUT)
    print(f'Saved: {OUT}')


if __name__ == '__main__':
    main()
